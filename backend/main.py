from fastapi import FastAPI, File, UploadFile, HTTPException, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session
import os
import shutil
import logging
import re
from pathlib import Path
import PyPDF2
from typing import List, Dict, Optional, Any
from collections import Counter
from dotenv import load_dotenv
import time
from datetime import datetime

# Importar base de datos
from database import get_db, init_db
import db_services as db_svc

# Importar cache, FTS y analytics
import cache_manager as cache
import fts_search as fts
import analytics as analytics_module

# Importar traductor con IA
import ai_translator

# Importar traductor
import translator

# Cargar variables de entorno
load_dotenv()

# Configurar logging para reducir spam
if os.getenv("FILTER_404_LOGS", "False").lower() == "true":
    logging.getLogger("uvicorn.access").setLevel(logging.ERROR)

# Configuración desde variables de entorno
HOST = os.getenv("HOST", "0.0.0.0")
PORT = int(os.getenv("PORT", "8000"))
DEBUG = os.getenv("DEBUG", "True").lower() == "true"
RELOAD = os.getenv("RELOAD", "True").lower() == "true"
CORS_ORIGINS = os.getenv("CORS_ORIGINS", "http://localhost:3000").split(",")
MAX_FILE_SIZE = int(os.getenv("MAX_FILE_SIZE", "50")) * 1024 * 1024  # MB a bytes
UPLOAD_FOLDER = os.getenv("UPLOAD_FOLDER", "pdfs")
RESULTS_FOLDER = os.getenv("RESULTS_FOLDER", "results")

app = FastAPI(
    title="PDF Query API", 
    version="1.0.0",
    debug=DEBUG
)

# Configurar CORS
# En desarrollo, permite orígenes específicos o usa allow_origin_regex para red local
app.add_middleware(
    CORSMiddleware,
    allow_origins=CORS_ORIGINS if not DEBUG else ["*"],  # En DEBUG permite todos los orígenes
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    allow_headers=["*"],
    expose_headers=["*"],
    allow_origin_regex=r"http://(localhost|127\.0\.0\.1|192\.168\.\d+\.\d+):\d+" if DEBUG else None
)

# Directorios
UPLOAD_DIR = Path(UPLOAD_FOLDER)
RESULTS_DIR = Path(RESULTS_FOLDER)

# Crear directorios si no existen
UPLOAD_DIR.mkdir(exist_ok=True)
RESULTS_DIR.mkdir(exist_ok=True)

# ========== Modelos Pydantic ==========

class QueryRequest(BaseModel):
    question: str
    filename: str

class MultiQueryRequest(BaseModel):
    question: str
    filenames: List[str]
    search_all: Optional[bool] = False

# ========== Funciones de análisis de texto ==========

def extract_pdf_text(file_path: Path) -> str:
    """Extraer todo el texto de un PDF"""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    except Exception as e:
        raise Exception(f"Error extrayendo texto del PDF: {str(e)}")

def extract_pdf_text_by_pages(file_path: Path) -> Dict[int, str]:
    """Extraer texto de un PDF separado por páginas"""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            pages_text = {}
            for page_num, page in enumerate(pdf_reader.pages, start=1):
                pages_text[page_num] = page.extract_text()
        return pages_text
    except Exception as e:
        raise Exception(f"Error extrayendo texto del PDF: {str(e)}")

def search_in_text(text: str, keywords: List[str]) -> List[Dict[str, str]]:
    """Buscar palabras clave en el texto y retornar contexto"""
    results = []
    lines = text.split('\n')
    
    for keyword in keywords:
        pattern = re.compile(re.escape(keyword), re.IGNORECASE)
        for i, line in enumerate(lines):
            if pattern.search(line):
                # Obtener contexto (línea anterior, actual y siguiente)
                context_start = max(0, i - 1)
                context_end = min(len(lines), i + 2)
                context = ' '.join(lines[context_start:context_end])
                
                results.append({
                    "keyword": keyword,
                    "line_number": i + 1,
                    "context": context.strip()
                })
    
    return results

def search_in_pages(pages_text: Dict[int, str], keywords: List[str]) -> List[Dict]:
    """Buscar palabras clave en páginas específicas del PDF"""
    results = []
    
    for page_num, page_text in pages_text.items():
        lines = page_text.split('\n')
        
        for keyword in keywords:
            pattern = re.compile(re.escape(keyword), re.IGNORECASE)
            for i, line in enumerate(lines):
                if pattern.search(line):
                    # Obtener contexto
                    context_start = max(0, i - 1)
                    context_end = min(len(lines), i + 2)
                    context = ' '.join(lines[context_start:context_end])
                    
                    results.append({
                        "keyword": keyword,
                        "page": page_num,
                        "line_in_page": i + 1,
                        "context": context.strip()[:300]  # Limitar contexto
                    })
    
    return results

def analyze_question(question: str) -> Dict:
    """Analizar la pregunta para extraer palabras clave y tipo de consulta"""
    question_lower = question.lower()
    
    # Identificar tipo de pregunta
    question_type = "general"
    if any(word in question_lower for word in ["qué", "que", "what"]):
        question_type = "definition"
    elif any(word in question_lower for word in ["cómo", "como", "how"]):
        question_type = "process"
    elif any(word in question_lower for word in ["cuándo", "cuando", "when"]):
        question_type = "temporal"
    elif any(word in question_lower for word in ["dónde", "donde", "where"]):
        question_type = "location"
    elif any(word in question_lower for word in ["por qué", "porque", "why"]):
        question_type = "reason"
    elif any(word in question_lower for word in ["cuánto", "cuanto", "how much", "how many"]):
        question_type = "quantity"
    
    # Extraer palabras clave (remover palabras comunes)
    stop_words = {
        'el', 'la', 'los', 'las', 'un', 'una', 'unos', 'unas',
        'de', 'del', 'al', 'a', 'en', 'con', 'por', 'para',
        'qué', 'que', 'cómo', 'como', 'cuándo', 'cuando',
        'dónde', 'donde', 'por', 'es', 'está', 'son', 'están',
        'the', 'a', 'an', 'in', 'on', 'at', 'to', 'for',
        'what', 'how', 'when', 'where', 'why', 'is', 'are'
    }
    
    words = re.findall(r'\w+', question_lower)
    keywords = [word for word in words if len(word) > 3 and word not in stop_words]
    
    return {
        "type": question_type,
        "keywords": keywords[:5]  # Limitar a 5 palabras clave principales
    }

def generate_answer(question: str, pdf_text: str, filename: str) -> str:
    """Generar respuesta basada en el contenido del PDF"""
    # Analizar la pregunta
    analysis = analyze_question(question)
    keywords = analysis["keywords"]
    
    if not keywords:
        return "No pude identificar palabras clave en tu pregunta. Por favor, intenta ser más específico."
    
    # Buscar información relevante
    search_results = search_in_text(pdf_text, keywords)
    
    if not search_results:
        return f"No encontré información relacionada con '{', '.join(keywords)}' en el documento {filename}."
    
    # Construir respuesta basada en los resultados
    answer_parts = [f"Basándome en el documento '{filename}', encontré lo siguiente:\n"]
    
    # Agrupar resultados por keyword
    keyword_contexts = {}
    for result in search_results[:5]:  # Limitar a 5 resultados más relevantes
        keyword = result["keyword"]
        if keyword not in keyword_contexts:
            keyword_contexts[keyword] = []
        keyword_contexts[keyword].append(result["context"])
    
    # Construir respuesta coherente
    for keyword, contexts in keyword_contexts.items():
        answer_parts.append(f"\n📌 Respecto a '{keyword}':")
        for i, context in enumerate(contexts[:2], 1):  # Máximo 2 contextos por keyword
            # Limitar longitud del contexto
            if len(context) > 200:
                context = context[:200] + "..."
            answer_parts.append(f"  {i}. {context}")
    
    # Agregar estadísticas
    total_matches = len(search_results)
    answer_parts.append(f"\n\n💡 Encontré {total_matches} referencia(s) relacionada(s) con tu pregunta.")
    
    return "\n".join(answer_parts)

def generate_answer_with_pages(question: str, file_path: Path, filename: str) -> Dict:
    """Generar respuesta con ubicaciones de página"""
    # Extraer texto por páginas
    pages_text = extract_pdf_text_by_pages(file_path)
    
    # Analizar la pregunta
    analysis = analyze_question(question)
    keywords = analysis["keywords"]
    
    if not keywords:
        return {
            "answer": "No pude identificar palabras clave en tu pregunta. Por favor, intenta ser más específico.",
            "keywords": [],
            "locations": [],
            "total_matches": 0
        }
    
    # Buscar en páginas específicas
    search_results = search_in_pages(pages_text, keywords)
    
    if not search_results:
        return {
            "answer": f"No encontré información relacionada con '{', '.join(keywords)}' en el documento {filename}.",
            "keywords": keywords,
            "locations": [],
            "total_matches": 0
        }
    
    # Construir respuesta
    answer_parts = [f"📄 Basándome en el documento '{filename}', encontré lo siguiente:\n"]
    
    # Agrupar por página
    pages_found = {}
    for result in search_results:
        page = result["page"]
        if page not in pages_found:
            pages_found[page] = []
        pages_found[page].append(result)
    
    # Construir respuesta por página
    locations = []
    for page, results in sorted(pages_found.items())[:5]:  # Máximo 5 páginas
        answer_parts.append(f"\n📍 **Página {page}:**")
        
        # Mostrar contextos únicos
        contexts_shown = set()
        for result in results[:2]:  # Máximo 2 resultados por página
            context = result["context"]
            if context not in contexts_shown:
                contexts_shown.add(context)
                if len(context) > 200:
                    context = context[:200] + "..."
                answer_parts.append(f"   • {context}")
                
        locations.append({
            "page": page,
            "keywords": [r["keyword"] for r in results],
            "preview": results[0]["context"][:150] + "..."
        })
    
    # Estadísticas
    total_pages = len(pages_found)
    total_matches = len(search_results)
    answer_parts.append(f"\n\n📊 **Resumen:**")
    answer_parts.append(f"• Encontré {total_matches} coincidencia(s) en {total_pages} página(s)")
    answer_parts.append(f"• Palabras clave buscadas: {', '.join(keywords)}")
    
    return {
        "answer": "\n".join(answer_parts),
        "keywords": keywords,
        "locations": locations,
        "total_matches": total_matches,
        "pages_found": list(pages_found.keys())
    }

def search_multiple_pdfs(question: str, filenames: List[str]) -> Dict:
    """Buscar en múltiples PDFs y agregar resultados"""
    all_results = []
    total_matches = 0
    documents_found = 0
    
    # Analizar la pregunta una sola vez
    analysis = analyze_question(question)
    keywords = analysis["keywords"]
    
    if not keywords:
        return {
            "answer": "No pude identificar palabras clave en tu pregunta. Por favor, intenta ser más específico.",
            "keywords": [],
            "results": [],
            "total_matches": 0,
            "documents_found": 0,
            "comparison": {}
        }
    
    # Buscar en cada PDF
    for filename in filenames:
        file_path = UPLOAD_DIR / filename
        
        if not file_path.exists():
            continue
            
        try:
            # Extraer texto por páginas
            pages_text = extract_pdf_text_by_pages(file_path)
            
            # Buscar en páginas
            search_results = search_in_pages(pages_text, keywords)
            
            if search_results:
                documents_found += 1
                doc_matches = len(search_results)
                total_matches += doc_matches
                
                # Agrupar por página
                pages_found = {}
                for result in search_results:
                    page = result["page"]
                    if page not in pages_found:
                        pages_found[page] = []
                    pages_found[page].append(result)
                
                # Construir ubicaciones
                locations = []
                for page, results in sorted(pages_found.items())[:3]:  # Máximo 3 páginas por documento
                    locations.append({
                        "page": page,
                        "keywords": [r["keyword"] for r in results],
                        "preview": results[0]["context"][:150] + "..."
                    })
                
                all_results.append({
                    "filename": filename,
                    "matches": doc_matches,
                    "pages_found": list(pages_found.keys()),
                    "locations": locations,
                    "total_pages": len(pages_text)
                })
        
        except Exception as e:
            # Si hay error con un PDF, continuar con los demás
            continue
    
    # Si no se encontró nada
    if not all_results:
        return {
            "answer": f"No encontré información relacionada con '{', '.join(keywords)}' en los {len(filenames)} documento(s) seleccionados.",
            "keywords": keywords,
            "results": [],
            "total_matches": 0,
            "documents_found": 0,
            "comparison": {}
        }
    
    # Construir respuesta comparativa
    answer_parts = [f"🔍 **Búsqueda en {len(filenames)} documento(s)**\n"]
    answer_parts.append(f"📊 **Resultados:** Encontré información en **{documents_found}** de {len(filenames)} documentos.\n")
    
    # Ordenar documentos por número de coincidencias
    all_results.sort(key=lambda x: x["matches"], reverse=True)
    
    # Mostrar resultados por documento
    for idx, doc_result in enumerate(all_results[:5], 1):  # Máximo 5 documentos
        answer_parts.append(f"\n📄 **{idx}. {doc_result['filename']}**")
        answer_parts.append(f"   • {doc_result['matches']} coincidencia(s) en {len(doc_result['pages_found'])} página(s)")
        answer_parts.append(f"   • Páginas: {', '.join(map(str, doc_result['pages_found'][:5]))}")
        
        # Mostrar preview de la primera ubicación
        if doc_result['locations']:
            first_location = doc_result['locations'][0]
            preview = first_location['preview'][:120] + "..."
            answer_parts.append(f"   • Vista previa (pág. {first_location['page']}): \"{preview}\"")
    
    # Resumen comparativo
    answer_parts.append(f"\n\n💡 **Conclusión:**")
    answer_parts.append(f"• Total de {total_matches} coincidencia(s) encontradas")
    answer_parts.append(f"• {documents_found} documento(s) contienen información relevante")
    answer_parts.append(f"• Palabras clave buscadas: {', '.join(keywords)}")
    
    # Comparación de documentos
    comparison = {
        "most_relevant": all_results[0]["filename"] if all_results else None,
        "documents_with_results": documents_found,
        "documents_without_results": len(filenames) - documents_found,
        "average_matches_per_doc": round(total_matches / documents_found, 2) if documents_found > 0 else 0
    }
    
    return {
        "answer": "\n".join(answer_parts),
        "keywords": keywords,
        "results": all_results,
        "total_matches": total_matches,
        "documents_found": documents_found,
        "comparison": comparison
    }

def generate_summary(text: str, max_sentences: int = 5) -> str:
    """Generar un resumen del texto"""
    # Dividir en oraciones
    sentences = re.split(r'[.!?]+', text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 20]
    
    if not sentences:
        return "No se pudo generar un resumen del documento."
    
    # Tomar las primeras oraciones significativas
    summary_sentences = sentences[:max_sentences]
    
    # Calcular estadísticas básicas
    total_words = len(text.split())
    total_chars = len(text)
    total_lines = len(text.split('\n'))
    
    summary = "📋 RESUMEN DEL DOCUMENTO\n\n"
    summary += "\n".join(f"• {s}." for s in summary_sentences)
    summary += f"\n\n📊 ESTADÍSTICAS:\n"
    summary += f"• Total de palabras: {total_words:,}\n"
    summary += f"• Total de caracteres: {total_chars:,}\n"
    summary += f"• Total de líneas: {total_lines:,}"
    
    return summary

def get_word_frequency(text: str, top_n: int = 20) -> List[Dict]:
    """Obtener las palabras más frecuentes"""
    # Palabras a ignorar (stop words en español e inglés)
    stop_words = {
        'el', 'la', 'los', 'las', 'un', 'una', 'unos', 'unas',
        'de', 'del', 'al', 'a', 'en', 'con', 'por', 'para', 'su', 'sus',
        'este', 'esta', 'estos', 'estas', 'ese', 'esa', 'esos', 'esas',
        'y', 'o', 'pero', 'si', 'no', 'ni', 'que', 'como', 'cuando',
        'donde', 'quien', 'cual', 'muy', 'más', 'menos', 'tan', 'tanto',
        'es', 'son', 'está', 'están', 'ser', 'estar', 'ha', 'han',
        'the', 'a', 'an', 'in', 'on', 'at', 'to', 'for', 'of', 'and',
        'or', 'but', 'if', 'is', 'are', 'was', 'were', 'be', 'been',
        'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would',
        'can', 'could', 'may', 'might', 'must', 'shall', 'should'
    }
    
    # Extraer palabras
    words = re.findall(r'\b[a-záéíóúñA-ZÁÉÍÓÚÑ]{4,}\b', text.lower())
    
    # Filtrar stop words
    filtered_words = [word for word in words if word not in stop_words]
    
    # Contar frecuencias
    word_counts = Counter(filtered_words)
    
    # Obtener las top N
    top_words = word_counts.most_common(top_n)
    
    return [{"word": word, "count": count} for word, count in top_words]

def translate_text(text: str, target_lang: str = "en") -> str:
    """Traducción básica de términos comunes (placeholder para API de traducción)"""
    # Esta es una implementación básica
    # En producción, usar Google Translate API, DeepL, etc.
    
    translations = {
        "es_to_en": {
            "documento": "document",
            "archivo": "file",
            "página": "page",
            "texto": "text",
            "contenido": "content",
            "información": "information",
            "datos": "data",
            "análisis": "analysis",
            "resumen": "summary",
            "estadísticas": "statistics"
        },
        "en_to_es": {
            "document": "documento",
            "file": "archivo",
            "page": "página",
            "text": "texto",
            "content": "contenido",
            "information": "información",
            "data": "datos",
            "analysis": "análisis",
            "summary": "resumen",
            "statistics": "estadísticas"
        }
    }
    
    # Nota: Esta es una implementación muy básica
    # Para traducción real, integrar con API de traducción
    return f"⚠️ Traducción completa requiere API externa. Texto original:\n\n{text[:500]}..."

# ========== Endpoints de la API ==========

@app.get("/")
async def root():
    return {
        "message": "PDF Query API está funcionando",
        "version": "1.0.0",
        "status": "online",
        "upload_folder": str(UPLOAD_DIR),
        "max_file_size_mb": MAX_FILE_SIZE // (1024*1024)
    }

@app.get("/health")
async def health_check():
    """Endpoint de verificación de salud"""
    return {
        "status": "healthy",
        "timestamp": str(Path.cwd()),
        "upload_dir_exists": UPLOAD_DIR.exists(),
        "results_dir_exists": RESULTS_DIR.exists()
    }

@app.post("/upload-pdf")
async def upload_pdf(file: UploadFile = File(...), db: Session = Depends(get_db)):
    """Subir un archivo PDF"""
    # Validar que se proporcionó un archivo
    if not file.filename:
        raise HTTPException(status_code=400, detail="No se proporcionó ningún archivo")
    
    # Validar extensión
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Solo se permiten archivos PDF")
    
    # Validar tamaño del archivo
    contents = await file.read()
    file_size = len(contents)
    if file_size > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=413, 
            detail=f"El archivo es demasiado grande. Tamaño máximo: {MAX_FILE_SIZE // (1024*1024)}MB"
        )
    
    # Guardar archivo
    file_path = UPLOAD_DIR / file.filename
    with open(file_path, "wb") as buffer:
        buffer.write(contents)
    
    try:
        # Extraer texto y metadata del PDF
        pages_text = extract_pdf_text_by_pages(file_path)
        full_text = "\n".join(pages_text.values())
        total_pages = len(pages_text)
        
        # Guardar en base de datos
        pdf_doc = db_svc.create_pdf_document(
            db=db,
            filename=file.filename,
            file_path=str(file_path),
            file_size=file_size,
            total_pages=total_pages,
            full_text=full_text,
            text_by_pages=pages_text
        )
        
        # Incrementar estadísticas
        db_svc.increment_upload_count(db)
        
        return {
            "message": f"Archivo {file.filename} subido correctamente",
            "file_path": str(file_path),
            "pdf_id": pdf_doc.id,
            "total_pages": total_pages,
            "word_count": pdf_doc.word_count
        }
    
    except Exception as e:
        # Si falla, eliminar archivo y registro de BD
        if file_path.exists():
            file_path.unlink()
        raise HTTPException(status_code=500, detail=f"Error procesando PDF: {str(e)}")

@app.post("/extract-text/{filename}")
async def extract_text(filename: str):
    """Extraer texto de un PDF"""
    file_path = UPLOAD_DIR / filename
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Archivo no encontrado")
    
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        
        # Guardar texto extraído
        text_file = RESULTS_DIR / f"{filename}_extracted.txt"
        with open(text_file, 'w', encoding='utf-8') as f:
            f.write(text)
        
        return {"message": "Texto extraído correctamente", "text": text[:500] + "..."}
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error al procesar PDF: {str(e)}")

@app.post("/query")
async def query_pdf(query: dict, db: Session = Depends(get_db)):
    """Realizar consulta sobre el contenido del PDF"""
    question = query.get("question", "")
    filename = query.get("filename", "")
    
    if not question or not filename:
        raise HTTPException(status_code=400, detail="Se requiere pregunta y nombre de archivo")
    
    # Verificar que el archivo existe
    file_path = UPLOAD_DIR / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail=f"Archivo {filename} no encontrado")
    
    try:
        # 🚀 INTENTAR RECUPERAR DEL CACHE
        cached_result = cache.get_cached_result(db, question, [filename], "single")
        if cached_result:
            print(f"✅ Cache HIT para query: {question[:50]}...")
            # Actualizar estadísticas (no contar tiempo de ejecución)
            db_svc.increment_query_count(db, 0.001, cached_result.get("total_matches", 0), "single")
            return {
                **cached_result,
                "cached": True,
                "cache_hit": True
            }
        
        print(f"❌ Cache MISS para query: {question[:50]}...")
        start_time = time.time()
        
        # Generar respuesta con ubicaciones de página
        result = generate_answer_with_pages(question, file_path, filename)
        
        execution_time = time.time() - start_time
        
        # Agregar información adicional
        result["question"] = question
        result["filename"] = filename
        result["cached"] = False
        
        # 💾 GUARDAR EN CACHE (TTL 24 horas por defecto)
        cache.cache_query_result(
            db=db,
            question=question,
            pdf_files=[filename],
            search_type="single",
            result=result,
            execution_time=execution_time,
            ttl_hours=24
        )
        
        # Guardar en historial
        db_svc.create_query_history(
            db=db,
            question=question,
            pdf_filename=filename,
            search_type="single",
            keywords_found=result.get("keywords", []),
            total_matches=result.get("total_matches", 0),
            documents_found=1,
            execution_time=execution_time,
            answer=result.get("answer", ""),
            results=result
        )
        
        # Actualizar acceso del PDF
        db_svc.update_pdf_access(db, filename)
        
        # Actualizar estadísticas
        db_svc.increment_query_count(db, execution_time, result.get("total_matches", 0), "single")
        if result.get("keywords"):
            db_svc.update_top_keywords(db, result["keywords"])
        
        return result
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error procesando consulta: {str(e)}")

@app.post("/query-multiple")
async def query_multiple_pdfs_endpoint(request: MultiQueryRequest, db: Session = Depends(get_db)):
    """Realizar consulta en múltiples PDFs"""
    question = request.question
    filenames = request.filenames
    search_all = request.search_all
    
    if not question:
        raise HTTPException(status_code=400, detail="Se requiere una pregunta")
    
    # Si search_all es True, buscar en todos los PDFs disponibles
    if search_all:
        filenames = [f.name for f in UPLOAD_DIR.glob("*.pdf")]
        
        if not filenames:
            raise HTTPException(status_code=404, detail="No se encontraron archivos PDF")
    
    if not filenames:
        raise HTTPException(status_code=400, detail="Se requiere al menos un archivo")
    
    try:
        search_type = "all" if search_all else "multiple"
        
        # 🚀 INTENTAR RECUPERAR DEL CACHE
        cached_result = cache.get_cached_result(db, question, filenames, search_type)
        if cached_result:
            print(f"✅ Cache HIT para query múltiple: {question[:50]}...")
            # Actualizar estadísticas (no contar tiempo de ejecución)
            db_svc.increment_query_count(db, 0.001, cached_result.get("total_matches", 0), search_type)
            return {
                **cached_result,
                "cached": True,
                "cache_hit": True
            }
        
        print(f"❌ Cache MISS para query múltiple: {question[:50]}...")
        start_time = time.time()
        
        # Realizar búsqueda en múltiples PDFs
        result = search_multiple_pdfs(question, filenames)
        
        execution_time = time.time() - start_time
        
        # Agregar información adicional
        result["question"] = question
        result["searched_files"] = filenames
        result["search_all"] = search_all
        result["cached"] = False
        
        # 💾 GUARDAR EN CACHE (TTL 12 horas para multi-búsquedas)
        cache.cache_query_result(
            db=db,
            question=question,
            pdf_files=filenames,
            search_type=search_type,
            result=result,
            execution_time=execution_time,
            ttl_hours=12  # Menor TTL para búsquedas múltiples
        )
        
        # Guardar en historial
        db_svc.create_query_history(
            db=db,
            question=question,
            multiple_pdfs=filenames,
            search_type=search_type,
            keywords_found=result.get("keywords", []),
            total_matches=result.get("total_matches", 0),
            documents_found=result.get("documents_found", 0),
            execution_time=execution_time,
            answer=result.get("answer", ""),
            results=result
        )
        
        # Actualizar acceso de cada PDF encontrado
        if result.get("results"):
            for doc_result in result["results"]:
                db_svc.update_pdf_access(db, doc_result["filename"])
        
        # Actualizar estadísticas
        db_svc.increment_query_count(db, execution_time, result.get("total_matches", 0), search_type)
        if result.get("keywords"):
            db_svc.update_top_keywords(db, result["keywords"])
        
        return result
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error procesando consulta múltiple: {str(e)}")

@app.get("/list-pdfs")
async def list_pdfs(db: Session = Depends(get_db)):
    """Listar todos los PDFs subidos con metadata"""
    pdfs_db = db_svc.get_all_pdfs(db)
    
    return {
        "pdfs": [pdf.filename for pdf in pdfs_db],
        "detailed": [
            {
                "filename": pdf.filename,
                "upload_date": pdf.upload_date.isoformat() if pdf.upload_date else None,
                "file_size": pdf.file_size,
                "total_pages": pdf.total_pages,
                "access_count": pdf.access_count,
                "tags": pdf.tags or [],
                "category": pdf.category
            }
            for pdf in pdfs_db
        ]
    }

@app.get("/view-pdf/{filename}")
async def view_pdf(filename: str, page: int = 1):
    """Servir PDF para visualización en el navegador"""
    from fastapi.responses import FileResponse
    
    file_path = UPLOAD_DIR / filename
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail=f"Archivo {filename} no encontrado")
    
    # Devolver el archivo PDF con headers apropiados
    return FileResponse(
        path=file_path,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f"inline; filename={filename}",
            "Access-Control-Expose-Headers": "Content-Disposition"
        }
    )

@app.post("/analyze/{filename}")
async def analyze_pdf(filename: str, analysis_type: str = "summary"):
    """Realizar análisis avanzado del PDF"""
    file_path = UPLOAD_DIR / filename
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail=f"Archivo {filename} no encontrado")
    
    try:
        # Extraer texto del PDF
        pdf_text = extract_pdf_text(file_path)
        
        if not pdf_text.strip():
            return {
                "error": f"El archivo {filename} no contiene texto extraíble.",
                "analysis_type": analysis_type
            }
        
        result: Dict[str, Any] = {
            "filename": filename,
            "analysis_type": analysis_type
        }
        
        if analysis_type == "summary":
            result["summary"] = generate_summary(pdf_text)
            
        elif analysis_type == "word_frequency":
            result["word_frequency"] = get_word_frequency(pdf_text, top_n=20)
            result["total_unique_words"] = len(set(pdf_text.lower().split()))
            
        elif analysis_type == "statistics":
            words = pdf_text.split()
            lines = pdf_text.split('\n')
            result["statistics"] = {
                "total_words": len(words),
                "total_characters": len(pdf_text),
                "total_characters_no_spaces": len(pdf_text.replace(" ", "")),
                "total_lines": len(lines),
                "total_pages": "N/A",  # PyPDF2 puede obtener esto
                "average_word_length": sum(len(word) for word in words) / len(words) if words else 0
            }
            
        elif analysis_type == "translate":
            # Traducción básica (placeholder)
            result["translated"] = translate_text(pdf_text[:1000])
            result["note"] = "Traducción completa requiere integración con API de traducción"
            
        else:
            raise HTTPException(status_code=400, detail=f"Tipo de análisis no soportado: {analysis_type}")
        
        return result
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error analizando PDF: {str(e)}")

@app.post("/batch-analyze/{filename}")
async def batch_analyze_pdf(filename: str):
    """Realizar todos los análisis de una vez"""
    file_path = UPLOAD_DIR / filename
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail=f"Archivo {filename} no encontrado")
    
    try:
        pdf_text = extract_pdf_text(file_path)
        
        if not pdf_text.strip():
            return {"error": f"El archivo {filename} no contiene texto extraíble."}
        
        words = pdf_text.split()
        
        return {
            "filename": filename,
            "summary": generate_summary(pdf_text),
            "word_frequency": get_word_frequency(pdf_text, top_n=15),
            "statistics": {
                "total_words": len(words),
                "total_characters": len(pdf_text),
                "total_lines": len(pdf_text.split('\n')),
                "unique_words": len(set(word.lower() for word in words)),
                "average_word_length": round(sum(len(word) for word in words) / len(words) if words else 0, 2)
            }
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error en análisis batch: {str(e)}")

# ========== ENDPOINTS DE BASE DE DATOS ==========

@app.get("/api/history")
async def get_query_history(limit: int = 20, db: Session = Depends(get_db)):
    """Obtener historial de consultas recientes"""
    queries = db_svc.get_recent_queries(db, limit=limit)
    
    return {
        "total": len(queries),
        "queries": [
            {
                "id": q.id,
                "question": q.question,
                "pdf_filename": q.pdf_filename,
                "search_type": q.search_type,
                "total_matches": q.total_matches,
                "documents_found": q.documents_found,
                "execution_time": q.execution_time,
                "query_date": q.query_date.isoformat() if q.query_date else None
            }
            for q in queries
        ]
    }

@app.get("/api/statistics")
async def get_statistics(days: int = 7, db: Session = Depends(get_db)):
    """Obtener estadísticas de uso"""
    stats = db_svc.get_statistics_summary(db, days=days)
    return stats

@app.get("/api/dashboard")
async def get_dashboard(db: Session = Depends(get_db)):
    """Obtener datos para dashboard"""
    dashboard_data = db_svc.get_dashboard_data(db)
    return dashboard_data

@app.get("/api/pdf/{filename}/stats")
async def get_pdf_stats(filename: str, db: Session = Depends(get_db)):
    """Obtener estadísticas de un PDF específico"""
    stats = db_svc.get_pdf_statistics(db, filename)
    
    if not stats:
        raise HTTPException(status_code=404, detail=f"PDF {filename} no encontrado en base de datos")
    
    return stats

@app.post("/api/pdf/{filename}/tags")
async def add_pdf_tags(filename: str, tags: List[str], db: Session = Depends(get_db)):
    """Agregar tags a un PDF"""
    db_svc.add_pdf_tags(db, filename, tags)
    return {"message": f"Tags agregados a {filename}", "tags": tags}

@app.post("/api/pdf/{filename}/category")
async def set_pdf_category(filename: str, category: str, db: Session = Depends(get_db)):
    """Establecer categoría de un PDF"""
    db_svc.set_pdf_category(db, filename, category)
    return {"message": f"Categoría '{category}' establecida para {filename}"}

@app.delete("/api/pdf/{filename}")
async def delete_pdf(filename: str, db: Session = Depends(get_db)):
    """Eliminar un PDF"""
    # Eliminar archivo físico
    file_path = UPLOAD_DIR / filename
    if file_path.exists():
        file_path.unlink()
    
    # Invalidar cache del PDF antes de eliminarlo
    invalidated_count = cache.invalidate_cache_for_pdf(db, filename)
    print(f"🗑️ Cache invalidado para {filename}: {invalidated_count} entradas")
    
    # Eliminar del índice FTS
    try:
        pdf = db_svc.get_pdf_by_filename(db, filename)
        if pdf:
            fts.remove_pdf_from_fts(db, pdf.id)
            print(f"🗑️ Eliminado del índice FTS: {filename}")
    except Exception as e:
        print(f"⚠️ Error eliminando de FTS: {e}")
    
    # Eliminar de base de datos
    deleted = db_svc.delete_pdf_document(db, filename)
    
    if deleted:
        return {
            "message": f"PDF {filename} eliminado correctamente",
            "cache_invalidated": invalidated_count,
            "fts_removed": True
        }
    else:
        raise HTTPException(status_code=404, detail=f"PDF {filename} no encontrado")

@app.get("/api/popular-queries")
async def get_popular_queries(limit: int = 10, db: Session = Depends(get_db)):
    """Obtener consultas más populares"""
    popular = db_svc.get_popular_queries(db, limit=limit)
    return {"popular_queries": popular}


# ============================================================
# ENDPOINTS DE CACHE
# ============================================================

@app.get("/api/cache/stats")
async def get_cache_stats(db: Session = Depends(get_db)):
    """Obtener estadísticas del cache"""
    try:
        stats = cache.get_cache_statistics(db)
        return stats
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error obteniendo estadísticas de cache: {str(e)}")


@app.post("/api/cache/clear")
async def clear_cache(expired_only: bool = True, db: Session = Depends(get_db)):
    """Limpiar cache (solo expirados o todo)"""
    try:
        if expired_only:
            removed = cache.clear_expired_cache(db)
            return {"message": f"Cache expirado limpiado", "removed_entries": removed}
        else:
            # Limpiar todo (marcar como inválido)
            from cache_manager import QueryCache
            all_cache = db.query(QueryCache).all()
            for entry in all_cache:
                entry.is_valid = False
            db.commit()
            return {"message": "Todo el cache ha sido invalidado", "removed_entries": len(all_cache)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error limpiando cache: {str(e)}")


@app.post("/api/cache/cleanup")
async def smart_cache_cleanup(max_entries: int = 1000, min_hits: int = 2, db: Session = Depends(get_db)):
    """Limpieza inteligente del cache"""
    try:
        result = cache.smart_cache_cleanup(db, max_entries=max_entries, min_hits=min_hits)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error en limpieza inteligente: {str(e)}")


# ============================================================
# ENDPOINTS DE FULL-TEXT SEARCH (FTS)
# ============================================================

@app.post("/api/fts/init")
async def initialize_fts(db: Session = Depends(get_db)):
    """Inicializar tablas FTS5"""
    try:
        fts.init_fts_tables(db)
        return {"message": "Tablas FTS5 inicializadas correctamente"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error inicializando FTS: {str(e)}")


@app.post("/api/fts/index/{filename}")
async def index_pdf_fts(filename: str, db: Session = Depends(get_db)):
    """Indexar un PDF específico en FTS"""
    try:
        # Obtener PDF de base de datos
        pdf = db_svc.get_pdf_by_filename(db, filename)
        if not pdf:
            raise HTTPException(status_code=404, detail=f"PDF {filename} no encontrado")
        
        if not pdf.text_by_pages:
            raise HTTPException(status_code=400, detail=f"PDF {filename} no tiene texto extraído")
        
        # Indexar
        fts.index_pdf_for_fts(db, pdf.id, pdf.filename, pdf.text_by_pages)
        return {"message": f"PDF {filename} indexado en FTS", "pages_indexed": len(pdf.text_by_pages)}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error indexando PDF: {str(e)}")


@app.post("/api/fts/rebuild")
async def rebuild_fts(db: Session = Depends(get_db)):
    """Reconstruir índice FTS completo"""
    try:
        result = fts.rebuild_fts_index(db)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error reconstruyendo FTS: {str(e)}")


@app.get("/api/fts/search")
async def search_fts(query: str, filenames: Optional[str] = None, limit: int = 50, db: Session = Depends(get_db)):
    """Búsqueda full-text ultrarrápida"""
    try:
        filename_list = filenames.split(",") if filenames else None
        results = fts.fts_search(db, query, filename_list, limit)
        return {
            "query": query,
            "total_results": len(results),
            "results": results
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error en búsqueda FTS: {str(e)}")


@app.get("/api/fts/stats")
async def get_fts_stats(db: Session = Depends(get_db)):
    """Estadísticas del índice FTS"""
    try:
        stats = fts.get_fts_statistics(db)
        return stats
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error obteniendo estadísticas FTS: {str(e)}")


# ============================================================
# ENDPOINTS DE ANALYTICS
# ============================================================

@app.get("/api/analytics/trending")
async def get_trending(days: int = 7, limit: int = 20, db: Session = Depends(get_db)):
    """Keywords en tendencia"""
    try:
        trending = analytics_module.get_trending_keywords(db, days=days, limit=limit)
        return {
            "period_days": days,
            "trending_keywords": trending
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error obteniendo tendencias: {str(e)}")


@app.get("/api/analytics/correlations")
async def get_correlations(days: int = 30, min_support: int = 2, db: Session = Depends(get_db)):
    """Correlaciones entre queries"""
    try:
        correlations = analytics_module.find_query_correlations(db, days=days, min_support=min_support)
        return {
            "period_days": days,
            "correlations": correlations
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error obteniendo correlaciones: {str(e)}")


@app.get("/api/analytics/similar-docs/{filename}")
async def get_similar_docs(filename: str, limit: int = 5, db: Session = Depends(get_db)):
    """Documentos similares a un PDF"""
    try:
        similar = analytics_module.find_similar_documents(db, filename, limit=limit)
        return {
            "target_document": filename,
            "similar_documents": similar
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error encontrando documentos similares: {str(e)}")


@app.get("/api/analytics/usage-patterns")
async def get_usage_patterns(days: int = 30, db: Session = Depends(get_db)):
    """Patrones de uso por hora y día"""
    try:
        patterns = analytics_module.get_usage_patterns_by_time(db, days=days)
        return patterns
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error obteniendo patrones de uso: {str(e)}")


@app.get("/api/analytics/pdf-trends")
async def get_pdf_trends(days: int = 30, db: Session = Depends(get_db)):
    """Tendencias de uso de PDFs"""
    try:
        trends = analytics_module.get_pdf_usage_trends(db, days=days)
        return {
            "period_days": days,
            "pdf_trends": trends
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error obteniendo tendencias de PDFs: {str(e)}")


@app.get("/api/analytics/performance")
async def get_performance_stats(days: int = 30, db: Session = Depends(get_db)):
    """Estadísticas de rendimiento"""
    try:
        stats = analytics_module.get_query_performance_stats(db, days=days)
        return stats
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error obteniendo estadísticas de rendimiento: {str(e)}")


@app.get("/api/analytics/user-patterns")
async def get_user_patterns(days: int = 30, db: Session = Depends(get_db)):
    """Patrones de comportamiento del usuario"""
    try:
        patterns = analytics_module.detect_user_patterns(db, days=days)
        return patterns
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error detectando patrones de usuario: {str(e)}")


@app.get("/api/analytics/dashboard")
async def get_analytics_dashboard(days: int = 30, db: Session = Depends(get_db)):
    """Dashboard completo de analytics"""
    try:
        dashboard = analytics_module.get_complete_analytics_dashboard(db, days=days)
        return dashboard
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generando dashboard de analytics: {str(e)}")


# ============================================================
# ENDPOINTS DE TRADUCCIÓN (ALEMÁN ↔ INGLÉS)
# ============================================================

class TranslateRequest(BaseModel):
    text: str
    source_lang: str = "de"  # alemán por defecto
    target_lang: str = "en"  # inglés por defecto
    preserve_case: bool = True

class TranslatePdfRequest(BaseModel):
    filename: str
    source_lang: str = "de"  # alemán por defecto
    target_lang: str = "en"  # inglés por defecto
    pages: Optional[List[int]] = None
    save_translated: bool = True
    output_format: str = "txt"  # "txt", "pdf", o "docx"
    use_ai: bool = True  # Usar IA si está disponible

@app.post("/api/translate")
async def translate_text_endpoint(request: TranslateRequest):
    """
    Traducir texto del alemán al inglés (o viceversa)
    
    Ejemplo:
    ```
    {
        "text": "Wie viele Seiten hat das Dokument?",
        "source_lang": "de",
        "target_lang": "en"
    }
    ```
    """
    try:
        result = translator.translate_query(
            request.text, 
            source_lang=request.source_lang,
            target_lang=request.target_lang
        )
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error traduciendo: {str(e)}")


@app.get("/api/translate/word")
async def translate_word_endpoint(
    word: str,
    source_lang: str = "de",
    target_lang: str = "en"
):
    """
    Traducir una palabra individual
    
    Ejemplo: /api/translate/word?word=dokument&source_lang=de&target_lang=en
    """
    try:
        translation = translator.translate_word(word, source_lang, target_lang)
        
        if translation:
            return {
                "original": word,
                "translation": translation,
                "source_lang": source_lang,
                "target_lang": target_lang,
                "found": True
            }
        else:
            return {
                "original": word,
                "translation": None,
                "source_lang": source_lang,
                "target_lang": target_lang,
                "found": False,
                "message": "Palabra no encontrada en diccionario"
            }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error traduciendo palabra: {str(e)}")


@app.get("/api/translate/stats")
async def get_translation_stats():
    """Estadísticas del diccionario de traducción"""
    try:
        stats = translator.get_dictionary_stats()
        return stats
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error obteniendo estadísticas: {str(e)}")


@app.post("/api/translate/custom")
async def add_custom_translation(german: str, english: str):
    """
    Agregar traducción personalizada al diccionario
    
    Ejemplo: /api/translate/custom?german=beispiel&english=example
    """
    try:
        translator.add_custom_translation(german, english)
        return {
            "message": f"Traducción agregada: {german} → {english}",
            "german": german,
            "english": english
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error agregando traducción: {str(e)}")


@app.post("/api/query-translated")
async def query_pdf_translated(
    question: str,
    filenames: List[str] = [],
    search_all: bool = False,
    source_lang: str = "de",
    target_lang: str = "en",
    translate_result: bool = False,
    db: Session = Depends(get_db)
):
    """
    Hacer query con traducción automática, soporta búsqueda múltiple
    
    Ejemplo:
    ```
    POST /api/query-translated
    {
        "question": "Wie viele Seiten hat das Dokument?",
        "filenames": ["doc1.pdf", "doc2.pdf"],
        "search_all": false,
        "source_lang": "de",
        "target_lang": "en",
        "translate_result": false
    }
    ```
    """
    try:
        # 1. Traducir pregunta
        translation_result = translator.translate_query(question, source_lang, target_lang)
        question_translated = translation_result["translated"]
        
        # 2. Determinar qué PDFs buscar
        if search_all:
            # Buscar en todos
            all_results = db_svc.search_all_pdfs(db, question_translated)
            
            response = {
                "original_question": question,
                "translated_question": question_translated,
                "translation": {
                    "original": question,
                    "translated": question_translated,
                    "coverage": translation_result.get("coverage_percentage", 0)
                },
                "answer": all_results.get("answer", ""),
                "results": all_results.get("results", []),
                "total_matches": all_results.get("total_matches", 0),
                "documents_found": all_results.get("documents_found", 0),
                "keywords": all_results.get("keywords", []),
                "comparison": all_results.get("comparison", {})
            }
            
        elif len(filenames) > 1:
            # Búsqueda múltiple
            multi_results = db_svc.search_multiple_pdfs(db, question_translated, filenames)
            
            response = {
                "original_question": question,
                "translated_question": question_translated,
                "translation": {
                    "original": question,
                    "translated": question_translated,
                    "coverage": translation_result.get("coverage_percentage", 0)
                },
                "answer": multi_results.get("answer", ""),
                "results": multi_results.get("results", []),
                "total_matches": multi_results.get("total_matches", 0),
                "documents_found": multi_results.get("documents_found", 0),
                "keywords": multi_results.get("keywords", []),
                "comparison": multi_results.get("comparison", {})
            }
            
        else:
            # Búsqueda individual
            if not filenames or len(filenames) == 0:
                raise HTTPException(status_code=400, detail="Debe especificar al menos un archivo")
            
            filename = filenames[0]
            file_path = UPLOAD_DIR / filename
            if not file_path.exists():
                raise HTTPException(status_code=404, detail=f"Archivo {filename} no encontrado")
            
            # Intenta cache primero
            cached_result = cache.get_cached_result(db, question_translated, [filename], "single")
            if cached_result:
                query_result = cached_result
                query_result["cached"] = True
            else:
                # Procesa query
                start_time = time.time()
                query_result = generate_answer_with_pages(question_translated, file_path, filename)
                execution_time = time.time() - start_time
                
                # Guarda en cache
                cache.cache_query_result(db, question_translated, [filename], "single", 
                                        query_result, execution_time, ttl_hours=24)
                query_result["cached"] = False
            
            response = {
                "original_question": question,
                "translated_question": question_translated,
                "translation": {
                    "original": question,
                    "translated": question_translated,
                    "coverage": translation_result.get("coverage_percentage", 0)
                },
                "answer": query_result.get("answer", ""),
                "locations": query_result.get("locations", []),
                "pages_found": query_result.get("pages_found", []),
                "total_matches": query_result.get("total_matches", 0),
                "keywords": query_result.get("keywords", []),
                "cached": query_result.get("cached", False)
            }
        
        # 3. Opcionalmente traducir resultado de vuelta
        if translate_result and response.get("answer"):
            answer_translated_back = translator.translate_text(
                response["answer"], 
                source_lang=target_lang, 
                target_lang=source_lang
            )
            response["answer_translated_back"] = answer_translated_back
        
        return response
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error en query traducida: {str(e)}")


@app.post("/api/translate-pdf")
async def translate_pdf_content(request: TranslatePdfRequest):
    """
    Traducir el contenido completo de un PDF (o páginas específicas)
    
    Ejemplo:
    ```
    POST /api/translate-pdf
    {
        "filename": "manual.pdf",
        "source_lang": "de",
        "target_lang": "en",
        "pages": [1, 2, 3],  // opcional, si no se especifica traduce todo
        "save_translated": true  // guardar PDF traducido
    }
    ```
    
    Retorna:
    - pages_translated: número de páginas traducidas
    - statistics: estadísticas de traducción
    - translated_file: nombre del archivo traducido (si save_translated=true)
    """
    try:
        file_path = UPLOAD_DIR / request.filename
        if not file_path.exists():
            raise HTTPException(status_code=404, detail=f"Archivo {request.filename} no encontrado")
        
        # Extraer texto del PDF por páginas
        pdf_text_by_pages = extract_pdf_text_by_pages(file_path)
        
        if not pdf_text_by_pages:
            raise HTTPException(status_code=400, detail="No se pudo extraer texto del PDF")
        
        # Determinar qué páginas traducir
        pages_to_translate = request.pages if request.pages else list(pdf_text_by_pages.keys())
        
        # Traducir cada página
        original_pages = {}
        translated_pages = {}
        translation_stats = {
            "total_pages": len(pages_to_translate),
            "total_words_original": 0,
            "total_words_translated": 0,
            "average_coverage": 0,
            "pages_with_low_coverage": []
        }
        
        coverage_sum = 0
        
        for page_num in pages_to_translate:
            if page_num not in pdf_text_by_pages:
                continue
            
            original_text = pdf_text_by_pages[page_num]
            original_pages[page_num] = original_text
            
            # Mejorar traducción dividiendo en párrafos para mejor contexto
            if not original_text.strip():
                translated_pages[page_num] = original_text
                continue
                
            paragraphs = original_text.split('\n\n')
            translated_paragraphs = []
            page_coverage_sum = 0
            page_word_count_original = 0
            page_word_count_translated = 0
            
            for paragraph in paragraphs:
                if not paragraph.strip():
                    translated_paragraphs.append(paragraph)
                    continue
                
                # Traducir párrafo con IA o diccionario local
                if request.use_ai:
                    try:
                        translation_result = await ai_translator.translate_with_ai(paragraph.strip(), request.source_lang, request.target_lang)
                    except Exception as e:
                        # Fallback al diccionario local si falla la IA
                        translation_result = translator.translate_query(paragraph.strip(), request.source_lang, request.target_lang)
                else:
                    translation_result = translator.translate_query(paragraph.strip(), request.source_lang, request.target_lang)
                    
                translated_paragraph = translation_result["translated"]
                translated_paragraphs.append(translated_paragraph)
                
                # Acumular estadísticas del párrafo
                para_words_orig = len(paragraph.split())
                para_words_trans = len(translated_paragraph.split())
                para_coverage = translation_result.get("coverage_percentage", 0)
                
                page_word_count_original += para_words_orig
                page_word_count_translated += para_words_trans
                page_coverage_sum += para_coverage * para_words_orig  # Peso por palabras
            
            # Unir párrafos traducidos
            translated_text = '\n\n'.join(translated_paragraphs)
            translated_pages[page_num] = translated_text
            
            # Calcular cobertura promedio de la página (ponderada por palabras)
            page_coverage = page_coverage_sum / page_word_count_original if page_word_count_original > 0 else 0
            
            # Estadísticas globales
            translation_stats["total_words_original"] += page_word_count_original
            translation_stats["total_words_translated"] += page_word_count_translated
            coverage_sum += page_coverage
            
            if page_coverage < 70:
                translation_stats["pages_with_low_coverage"].append({
                    "page": page_num,
                    "coverage": round(page_coverage, 2)
                })
        
        # Calcular promedio de cobertura
        if len(pages_to_translate) > 0:
            translation_stats["average_coverage"] = round(coverage_sum / len(pages_to_translate), 2)
        
        # Guardar archivo traducido si se solicita
        translated_filename = None
        if request.save_translated:
            # Crear nombre de archivo traducido
            base_name = request.filename.rsplit('.', 1)[0]
            if request.output_format == "docx":
                file_extension = "docx"
            elif request.output_format == "pdf":
                file_extension = "pdf"
            else:
                file_extension = "txt"
                
            translated_filename = f"{base_name}_{request.source_lang}_to_{request.target_lang}.{file_extension}"
            translated_path = RESULTS_DIR / translated_filename
            
            if request.output_format == "docx":
                # Crear documento Word
                try:
                    from docx import Document
                    from docx.shared import Inches
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    
                    doc = Document()
                    
                    # Agregar título principal
                    title = doc.add_heading(f"PDF Traducido: {request.filename}", 0)
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Agregar metadata
                    doc.add_paragraph(f"Idioma: {request.source_lang.upper()} → {request.target_lang.upper()}")
                    doc.add_paragraph(f"Páginas: {len(translated_pages)}")
                    doc.add_paragraph(f"Cobertura promedio: {translation_stats['average_coverage']}%")
                    doc.add_paragraph(f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                    doc.add_paragraph()
                    
                    # Agregar contenido por páginas
                    for page_num in sorted(translated_pages.keys()):
                        # Título de página
                        page_title = doc.add_heading(f"Página {page_num}", level=1)
                        
                        # Contenido de la página
                        paragraphs = translated_pages[page_num].split('\n\n')
                        for paragraph_text in paragraphs:
                            if paragraph_text.strip():
                                doc.add_paragraph(paragraph_text.strip())
                        
                        # Espacio entre páginas
                        doc.add_page_break()
                    
                    doc.save(str(translated_path))
                    
                except ImportError:
                    # Si python-docx no está instalado, usar formato TXT
                    translated_filename = f"{base_name}_{request.source_lang}_to_{request.target_lang}.txt"
                    translated_path = RESULTS_DIR / translated_filename
                    request.output_format = "txt"
                    
            elif request.output_format == "pdf":
                # Crear PDF traducido manteniendo estructura
                try:
                    from reportlab.lib.pagesizes import A4
                    from reportlab.pdfgen import canvas
                    from reportlab.lib.styles import getSampleStyleSheet
                    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                    from reportlab.lib.units import inch
                    
                    doc = SimpleDocTemplate(str(translated_path), pagesize=A4)
                    styles = getSampleStyleSheet()
                    story = []
                    
                    for page_num in sorted(translated_pages.keys()):
                        # Agregar título de página
                        title = Paragraph(f"<b>Página {page_num}</b>", styles['Heading2'])
                        story.append(title)
                        story.append(Spacer(1, 0.2*inch))
                        
                        # Agregar contenido traducido
                        text = translated_pages[page_num].replace('\n', '<br/>')
                        content = Paragraph(text, styles['Normal'])
                        story.append(content)
                        story.append(Spacer(1, 0.3*inch))
                    
                    doc.build(story)
                except ImportError:
                    # Si reportlab no está instalado, usar formato TXT
                    translated_filename = f"{base_name}_{request.source_lang}_to_{request.target_lang}.txt"
                    translated_path = RESULTS_DIR / translated_filename
                    request.output_format = "txt"
            
            if request.output_format == "txt":
                # Guardar como TXT mejorado
                with open(translated_path, 'w', encoding='utf-8') as f:
                    f.write(f"# PDF TRADUCIDO: {request.filename}\n")
                    f.write(f"# Idioma: {request.source_lang.upper()} → {request.target_lang.upper()}\n")
                    f.write(f"# Páginas: {len(translated_pages)}\n")
                    f.write(f"# Cobertura promedio: {translation_stats['average_coverage']}%\n")
                    f.write(f"# Generado: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                    f.write(f"\n{'='*80}\n\n")
                    
                    for page_num in sorted(translated_pages.keys()):
                        f.write(f"{'='*60}\n")
                        f.write(f"PÁGINA {page_num}\n")
                        f.write(f"{'='*60}\n\n")
                        f.write(translated_pages[page_num])
                        f.write(f"\n\n{'='*60}\n\n")
        
        return {
            "filename": request.filename,
            "source_lang": request.source_lang,
            "target_lang": request.target_lang,
            "pages_translated": len(translated_pages),
            "statistics": translation_stats,
            "translated_file": translated_filename if request.save_translated else None,
            "download_url": f"/api/download-translated/{translated_filename}" if translated_filename else None
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error al traducir PDF: {str(e)}")


@app.get("/api/ai-info")
async def get_ai_translation_info():
    """
    Obtener información sobre disponibilidad de IA para traducción
    """
    try:
        info = ai_translator.get_ai_info()
        return {
            "ai_available": info["method"] in ["gemini", "openai"],
            "method": info["method"],
            "gemini_available": info["gemini_available"],
            "openai_available": info["openai_available"],
            "local_fallback": info["local_fallback"]
        }
    except Exception as e:
        return {
            "ai_available": False,
            "method": "local",
            "error": str(e)
        }


@app.get("/api/download-translated/{filename}")
async def download_translated_file(filename: str):
    """
    Descargar archivo traducido
    """
    file_path = RESULTS_DIR / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail=f"Archivo traducido {filename} no encontrado")
    
    return FileResponse(
        file_path,
        media_type="text/plain",
        filename=filename,
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


@app.get("/api/translated-files")
async def list_translated_files():
    """
    Listar todos los archivos traducidos disponibles
    """
    try:
        if not RESULTS_DIR.exists():
            return {"translated_files": []}
        
        translated_files = [
            {
                "filename": f.name,
                "size": f.stat().st_size,
                "created": f.stat().st_mtime,
                "download_url": f"/api/download-translated/{f.name}"
            }
            for f in RESULTS_DIR.glob("*_to_*.txt")
        ]
        
        return {
            "count": len(translated_files),
            "translated_files": translated_files
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error al listar archivos: {str(e)}")


# Catch-all para peticiones de hot-reload de React (evitar spam de 404)
@app.get("/{path:path}")
async def catch_react_routes(path: str):
    """Manejar rutas de React que no son API"""
    if path.endswith('.hot-update.json') or path.startswith('static/'):
        raise HTTPException(status_code=404, detail="React hot-reload file not found")
    
    # Para otras rutas, devolver info básica
    return {
        "message": "Esta es la API del backend", 
        "path_requested": path,
        "available_endpoints": ["/", "/docs", "/health", "/upload-pdf", "/list-pdfs"]
    }

if __name__ == "__main__":
    import uvicorn
    
    print(f"🚀 Iniciando servidor en http://{HOST}:{PORT}")
    print(f"📁 Directorio de PDFs: {UPLOAD_DIR.absolute()}")
    print(f"📊 Directorio de resultados: {RESULTS_DIR.absolute()}")
    print(f"🌐 CORS habilitado para: {', '.join(CORS_ORIGINS)}")
    
    if RELOAD:
        # Para reload, usar el string del módulo
        uvicorn.run(
            "main:app", 
            host=HOST, 
            port=PORT, 
            reload=True,
            log_level=os.getenv("LOG_LEVEL", "info").lower()
        )
    else:
        # Para producción, usar el objeto app directamente
        uvicorn.run(
            app, 
            host=HOST, 
            port=PORT, 
            log_level=os.getenv("LOG_LEVEL", "info").lower()
        )